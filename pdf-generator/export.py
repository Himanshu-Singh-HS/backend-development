r"""src/monolith/file_generator/drafting_export.py"""

# importing standard modules ==================================================
from typing import List, Dict, Any
import logging
import re
from io import BytesIO
 
# importing third-party modules ===============================================
import pdfkit
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests

 
# importing third-party modules  for pdf ===============================================
import time
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, PageBreak,Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY,TA_CENTER
import json
import os,requests
from io import BytesIO
from reportlab.lib.units import inch

# importing custom modules ====================================================
from monolith.drafting.exceptions import handle_errors
from monolith.drafting.service import time_decorator

# logging =====================================================================
logger: logging.Logger = logging.getLogger(__name__)

# document object ==============================================================
doc = Document()

#global list declare for generating pdf ===========================================================
elements=[]

@handle_errors
def extract_claims(claims_data: List[Dict[str, Any]]) -> Dict[str, List[str]]:
    """Extract independent, dependent claims and their figures for method and system claims."""
    sys_ind_claims, sys_dep_claims, sys_figs = [], [], []
    met_ind_claims, met_dep_claims, met_figs = [], [], []

    for claim in claims_data:
        claim_type = claim.get("claim_type")
        tag = claim.get("tag")
        text = claim.get("text", "")
        index = claim.get("index", "")
        latex_details = claim.get("generated_figures_data", {}) or {}
        latex_details = latex_details.get("latex_details", [])

        figs = [
            url
            for latex_detail in latex_details
            if latex_detail.get("images_urls", [])
            for url in latex_detail.get("images_urls", [])
        ]
        text = f"{index}. {text}"

        if claim_type == "method":
            if tag == "dependent":
                met_dep_claims.append(text)
            elif tag == "independent":
                met_ind_claims.append(text)
            if figs:
                met_figs.extend(figs)

        elif claim_type == "system":
            if tag == "dependent":
                sys_dep_claims.append(text)
            elif tag == "independent":
                sys_ind_claims.append(text)
            if figs:
                sys_figs.extend(figs)

    return {
        "sys_ind_claims": sys_ind_claims,
        "sys_dep_claims": sys_dep_claims,
        "sys_figs": sys_figs,
        "met_ind_claims": met_ind_claims,
        "met_dep_claims": met_dep_claims,
        "met_figs": met_figs,
    }


@handle_errors
def extract_description(description_data: Dict[str, Any]) -> Dict[str, List[str]]:
    """Extract method, system, and invention descriptions."""
    method_desc, system_desc, invention_desc = [], [], []

    for key, value in description_data.items():
        if key == "method_desc":
            method_desc.extend(value.get("text", "").split("\n\n"))
        elif key == "system_desc":
            system_desc.extend(value.get("text_list", []))
        elif key == "invention_desc":
            invention_desc.extend(value.get("text", "").split("\n\n"))

    return {
        "method_desc": method_desc,
        "system_desc": system_desc,
        "invention_desc": invention_desc,
    }


@handle_errors
def extract_list_of_figures(list_of_figures_data: Any) -> List[str]:
    """Extract list of figures."""
    if isinstance(list_of_figures_data, list):
        return list_of_figures_data
    return []


@handle_errors
def process_data(data: Dict[str, Any]) -> Dict[str, Any]:
    """Process the entire data set to extract claims, descriptions, and list of figures."""
    result = {
        "sys_ind_claims": [],
        "sys_dep_claims": [],
        "sys_figs": [],
        "met_ind_claims": [],
        "met_dep_claims": [],
        "met_figs": [],
        "method_desc": [],
        "system_desc": [],
        "invention_desc": [],
        "list_of_figures": [],
    }

    for key, value in data.items():
        if key == "description":
            desc_result = extract_description(value)
            result.update(desc_result)
        elif key == "claims":
            claims_result = extract_claims(value)
            result.update(claims_result)
        elif key == "list_of_figures":
            result["list_of_figures"] = extract_list_of_figures(value)
        elif isinstance(value, dict):
            result[key] = value.get("text", "").split("\n\n")
        else:
            result[key] = value.split("\n\n")

    return result


def generate_numbered_html_section(
    start_number: int, field_data_heading: str, field_data: list
) -> tuple:
    """
    Generates an HTML section with a heading and a list of items formatted with numbered paragraphs.
    Also returns the updated start number after processing all items.

    Args:
        start_number (int): The starting number for the list.
        field_data_heading (str): The heading text to be displayed.
        field_data (list): The list of items to be included in the HTML.

    Returns:
        tuple: A tuple containing:
            - str: The generated HTML section.
            - int: The updated start number.
    """
    heading = f"<p><b><span >{field_data_heading}</span></b></p>"
    if field_data_heading in {"ABSTRACT", "CLAIMS"}:
        heading = f'<p style="text-align:center;" ><b><span >{field_data_heading}</span></b></p>'

    if field_data_heading in {"FIGURES", "CLAIMS", "ABSTRACT"}:
        page_break = '<div style="page-break-before: always;"></div>'
        heading = page_break + heading

    if field_data_heading == "FIGURES":
        image_tags = "".join(
            f'<a href={fig} download="fig" > <img style="display: block; margin-left: auto; margin-right: auto; width: 60%;" src="{fig}" alt="Figure"></a>'
            for fig in field_data
        )
        html_section = heading + image_tags
        updated_start_number = start_number
    elif field_data_heading == "CLAIMS":

        def format_claims(field_data):
            formatted_claims = []
            for claim in field_data:
                split_claims = re.split(r",\s*and\s*|[;:]", claim)
                if split_claims:
                    formatted_claims.append(
                        f"<p>&nbsp;&nbsp; &nbsp; &nbsp; &nbsp;&nbsp; &nbsp;{split_claims[0]}</p>"
                    )
                    for claim_text in split_claims[1:]:
                        formatted_claims.append(f"<p>{claim_text}</p>")
            return "".join(formatted_claims)

        claim_body = f"""<div >
                {heading}
                <p><span
                        >What is claimed is: </span></p>
                   {format_claims(field_data)}
                
            </div>"""
        html_section = claim_body
        updated_start_number = start_number
    else:
        is_abstract = True if field_data_heading == "ABSTRACT" else False
        formatted_p_tags = [
            f'<p><b><span >{f"[{str(start_number + i).zfill(4)}]&nbsp;" if not is_abstract else "&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"}</span></b>'
            f"<b></b>"
            f"<span >{item}</span></p>"
            for i, item in enumerate(field_data)
        ]
        html_section = heading + "".join(formatted_p_tags)
        updated_start_number = start_number + len(field_data)

    return html_section, updated_start_number


def generate_final_html(data) -> str:
    """
    Combines all HTML sections into a final HTML document.

    Args:
        fields (list): A list of dictionaries where each dictionary contains:
            - 'heading': The heading text.
            - 'data': The list of items to be included in the HTML.

    Returns:
        str: The final HTML document.
    """
    keys_to_exclude = {
        "insert_timestamp",
        "last_update_timestamp",
        "_id",
        "user_id",
        "search_id",
        "components",
        "id",
        "status",
    }

    filtered_dict = {k: v for k, v in data.items() if k not in keys_to_exclude}

    processed_data = process_data(filtered_dict)
    css_styles = """
    <style>
         @page {
            size: A4;
            margin: 25.4mm;
        }
        body {
            font-family: 'Times New Roman', serif;
            font-size: 20px;
        }
   
        p{
            text-align:justify;
            line-height:1.8;
            
        }
        
    </style>
    """

    fields = [
        {
            "heading": "BACKGROUND",
            "data": processed_data.get("technical_field", "")
            + processed_data.get("background", ""),
        },
        {"heading": "BRIEF SUMMARY", "data": processed_data.get("summary", [])},
        {
            "heading": "BRIEF DESCRIPTION OF THE SEVERAL VIEWS OF THE DRAWINGS",
            "data": processed_data.get("list_of_figures", []),
        },
        {
            "heading": "DETAILED DESCRIPTION",
            "data": processed_data.get("method_desc", [])
            + processed_data.get("system_desc", [])
            + processed_data.get("invention_desc", []),
        },
        {
            "heading": "CLAIMS",
            "data": processed_data.get("sys_ind_claims", [])
            + processed_data.get("sys_dep_claims", [])
            + processed_data.get("met_ind_claims", [])
            + processed_data.get("met_dep_claims", []),
        },
        {"heading": "ABSTRACT", "data": processed_data.get("abstract", "")},
        {
            "heading": "FIGURES",
            "data": processed_data.get("met_figs", [])
            + processed_data.get("sys_figs", []),
        },
    ]
    start_number = 1
    html_sections = []

    for field in fields:
        html_section, start_number = generate_numbered_html_section(
            start_number, field["heading"], field["data"]
        )
        html_sections.append(html_section)

    title_html = f"""<p style="text-align:center;"><b><span>{processed_data.get('title', 'Document')[0].upper()}</span></b></p>"""
    final_html = f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><title></title>{css_styles}</head><body>"""
    final_html += title_html

    for i, section in enumerate(html_sections):
        final_html += section

    final_html += "</body></html>"

    return final_html

 

@time_decorator
def convert_html_to_pdf(html_content):
    """
    Convert HTML content to a PDF with custom settings using pdfkit.

    Parameters:
    - html_content: str, the HTML content to be converted to PDF.

    Returns:
    - The binary content of the generated PDF as a BytesIO object.
    """

    options = {
        "page-size": "A4",
        "footer-center": "[page]",
        "footer-font-size": "12",
        "margin-bottom": "20mm",
        "margin-left": "25.4mm",
        "margin-right": "20mm",
        "margin-top": "30mm",
        "encoding": "UTF-8",
        "header-right": "Attorney Docket No.                                 ",
        "header-spacing": "2",
        "header-font-size": "12",
        "header-font-name": "Times New Roman",
        "enable-local-file-access": "",  # Allows access to local files, such as CSS or JS
        "enable-javascript": "",  # Allows JavaScript to run in the HTML
    }
    return BytesIO(pdfkit.from_string(html_content, options=options))

def add_heading_with_color(text, level):
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        doc.add_paragraph()

def add_paragraph_with_numbering(text, counter):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"[{counter:04d}] ")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        section_run = paragraph.add_run(text)
        section_run.font.name = 'Times New Roman'
        section_run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        return counter + 1  # Increment counter
    
def convert_json_to_doc_buffer(Data: dict) -> BytesIO:
    """
    Converts a JSON Data to a Word document buffer.

    This function takes a dictionary (in JSON format) as input and generates
    a Word document (.docx).

    Args:
        Data (dict): A dictionary containing various sections of the document.
            Expected keys include:
            - "title": Contains a dictionary with the "text" field for the document title.
            - "technical_field": Dictionary with a "text" field for the technical field section.
            - "background": Dictionary with a "text" field containing background sections.
            - "summary": Dictionary with a "text" field for the summary section.
            - "list_of_figures": List of figure descriptions.
            - "description": Dictionary containing sections like "method_desc", "system_desc",
              and "invention_desc" with corresponding "text" or "text_list" fields.
            - "claims": List of claims, each with "text", "claim_type", and "generated_figures_data".
            - "abstract": Dictionary with a "text" field for the abstract section.
    
    Returns:
        BytesIO: A buffer containing the generated Word document in binary format.

    """
    # Title
    title_heading = doc.add_heading(Data["title"]["text"].upper(), level=1)
    title_run = title_heading.runs[0]
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(12)
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    counter = 1

    # Background attached with technical field
    add_heading_with_color('BACKGROUND', level=1)

    # Add technical field text with numbering
    counter = add_paragraph_with_numbering(
        Data["technical_field"]["text"], counter)

    # Background text
    background_text = Data["background"]["text"]
    sections = background_text.split('\n\n')
    for section in sections:
        counter = add_paragraph_with_numbering(section, counter)

    # Brief summary text
    add_heading_with_color('BRIEF SUMMARY', level=1)
    summary_text=Data["summary"]["text"]
    sections=summary_text.split('\n\n')
    for section in sections:
        counter = add_paragraph_with_numbering(section, counter)

    # List of figures
    add_heading_with_color(
        'BRIEF DESCRIPTION OF THE SEVERAL VIEWS OF THE DRAWINGS', level=1)
    for figure in Data["list_of_figures"]:
        counter = add_paragraph_with_numbering(figure, counter)

    # Detailed Description
    add_heading_with_color('DETAILED DESCRIPTION', level=1)

    # Method description
    method_desc_text = Data["description"]["method_desc"]["text"]
    method_desc_sections = method_desc_text.split('\n\n')
    for section in method_desc_sections:
        counter = add_paragraph_with_numbering(section, counter)

    # System description
    system_desc_text_list = Data["description"]["system_desc"]["text_list"]
    for section in system_desc_text_list:
        cleaned_text = section.replace('\n\n', '')
        counter = add_paragraph_with_numbering(cleaned_text, counter)

    # Invention Description
    invention_desc_text = Data["description"]["invention_desc"]["text"]
    invention_desc_sections = invention_desc_text.split('\n\n')
    for section in invention_desc_sections:
        counter = add_paragraph_with_numbering(section, counter)

    # Claims Description section
    doc.add_page_break()
    claims_heading = doc.add_heading('CLAIMS', level=1)
    claims_run = claims_heading.runs[0]
    claims_run.font.color.rgb = RGBColor(0, 0, 0)
    claims_run.font.name = 'Times New Roman'
    claims_run.font.size = Pt(12)
    claims_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Claims
    doc.add_paragraph("What is claimed is:")
    doc.add_paragraph()
    claim_indent = Pt(40)
    for index, claim in enumerate(Data['claims'], start=1):
        claim_parts = claim['text'].split('\n')
        for part in claim_parts:
            claim_paragraph = doc.add_paragraph()
            claim_paragraph.paragraph_format.first_line_indent = claim_indent 
            if part == claim_parts[0]:
                claim_run = claim_paragraph.add_run(f"{index}. ")
                claim_run.font.name = 'Times New Roman'
                claim_run.font.size = Pt(12)
            claim_text_run = claim_paragraph.add_run(part)
            claim_text_run.font.name = 'Times New Roman'
            claim_text_run.font.size = Pt(12)
        doc.add_paragraph()   
        
    
         
    
    # Abstract text
    doc.add_page_break()
    abstract_heading = doc.add_heading('ABSTRACT', level=1)
    doc.add_paragraph()
    abstract_run = abstract_heading.runs[0]
    abstract_run.font.color.rgb = RGBColor(0, 0, 0)
    abstract_run.font.name = 'Times New Roman'
    abstract_run.font.size = Pt(12)
    abstract_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    abstract_paragraph = doc.add_paragraph(Data["abstract"]["text"])

    for run in abstract_paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    abstract_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    abstract_paragraph.paragraph_format.first_line_indent = Pt(50)

    # Figures
    doc.add_page_break()
    add_heading_with_color('FIGURES', level=1)
    sorted_claims = sorted(
        Data["claims"], key=lambda x: x.get("claim_type") == "system")

    for claim in sorted_claims:
        if claim.get("generated_figures_data") and claim["generated_figures_data"].get("latex_details"):
            claim_type = claim.get("claim_type")
            for latex_detail in claim["generated_figures_data"]["latex_details"]:
                if "images_urls" in latex_detail:
                    for img_url in latex_detail["images_urls"]:
                        img_response = requests.get(img_url)
                        if img_response.status_code == 200:
                            image_stream = BytesIO(img_response.content)
                            image_paragraph = doc.add_paragraph()
                            run = image_paragraph.add_run()
                            run.add_picture(image_stream, width=Inches(4))
                            image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            print(
                                f"Failed to download the image from {img_url}")

    # Add "Attorney Docket No." to the header
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.add_paragraph()
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    header_paragraph.add_run("Attorney Docket No.")

    # Function to add page numbering in footer
    def add_page_numbering(doc):
        section = doc.sections[0]
        footer = section.footer
        footer_paragraph = footer.add_paragraph()
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = footer_paragraph.add_run()
        run._element.append(
            docx.oxml.parse_xml(
                r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
        )
        run._element.append(
            docx.oxml.parse_xml(
                r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE  \\* MERGEFORMAT </w:instrText>'
            )
        )
        run._element.append(
            docx.oxml.parse_xml(
                r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
        )

    add_page_numbering(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


#function for json to pdf 
# Create a custom paragraph style for justified alignment
style_sheet = getSampleStyleSheet()
justified_style = ParagraphStyle(
    'Justified',
    parent=style_sheet['BodyText'],
    fontName='Times-Roman',
    fontSize=12,
    leading=20,
    alignment=TA_JUSTIFY
)
centered_heading_style = ParagraphStyle(
    'CenteredHeading1',
    parent=style_sheet['Heading1'],
    alignment=TA_CENTER,
    fontSize=12,
    fontName='Times-Bold',  
)
heading2_style = ParagraphStyle(
    'Heading2',
    parent=style_sheet['Heading2'],
    # fontName='Times-Roman',  
    fontSize=12,    
    fontName='Times-Bold',       
)
title_style = ParagraphStyle(
    'Heading2',
    parent=style_sheet['Heading2'],
    # fontName='Times-Roman',  
    fontSize=12,    
    fontName='Times-Bold', 
    alignment=TA_CENTER      
)

def add_page_number_and_docket(canvas, doc, docket_number="Attorney Docket No. "):
    """This method is called to add the page number and docket number to each page."""
    page_num = canvas.getPageNumber()
    text = f"{page_num}"
    # Add the page number  
    canvas.saveState()
    canvas.setFont('Times-Roman', 12)
    canvas.drawCentredString(4.25 * inch, 0.5 * inch, text)
    # Add Attorney Docket No. 
    extra_space = 1 * inch   
    extra_space_down = 0.5 * inch
    canvas.setFont('Times-Roman', 12)
    canvas.drawRightString((7.5 * inch) - extra_space, (10.5 * inch-extra_space_down), docket_number)
    canvas.restoreState()

def download_image(img_url):
    """Downloads an image from a URL and returns it as a BytesIO object."""
    try:
        response = requests.get(img_url)
        if response.status_code == 200:
            return BytesIO(response.content)
        else:
            print(f"Failed to download the image from {img_url}")
    except Exception as e:
        print(f"Error downloading image: {e}")
    return None
 # Function to add justified paragraphs with numbering
def add_justified_paragraph_with_numbering(text, counter=0,  first_Line_Indent=0):
     
    modified_style = ParagraphStyle(
    'Justified',
    parent=style_sheet['BodyText'],
    fontName='Times-Roman',
    fontSize=12,
    leading=20,
    alignment=TA_JUSTIFY,
    firstLineIndent=first_Line_Indent)                                  
    numbered_text = f"<b>[{counter:04d}]</b> {text}"
    if counter==0:
        elements.append(Paragraph(text, modified_style))
    else:
        elements.append(Paragraph(numbered_text, modified_style))
        elements.append(Spacer(1, 12))  
        return counter + 1
    
    
  
def convert_json_to_pdf_buffer(Data: dict)-> BytesIO:
    """
    Converts a JSON dictionary into a PDF document and returns the document as a BytesIO buffer.

    Args:
        data (dict): The JSON data in dictionary format where each key-value pair will be added to the PDF.

    Returns:
        BytesIO: A BytesIO buffer containing the generated PDF. The buffer can be saved directly or sent as a response in a web application.
        
    # Use pdf_buffer to save the PDF or send as a response
    """
    buffer = BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=40)
    # Add title
    title_text = Data['title']['text'].upper()
    elements.append(Paragraph(title_text,title_style))
    elements.append(Spacer(1, 12)) 
        
    # Technical field text with numbering
    elements.append(Paragraph("BACKGROUND", heading2_style))
    counter = 1
    counter = add_justified_paragraph_with_numbering(Data["technical_field"]["text"], counter)

    # Background text
    background_text = Data["background"]["text"]
    sections = background_text.split('\n\n')
    for section in sections:
        counter = add_justified_paragraph_with_numbering(section, counter)

    # Brief summary text
    elements.append(Paragraph("BRIEF SUMMARY", heading2_style))
    elements.append(Spacer(1, 12))
    summary_text = Data["summary"]["text"]
    sections = summary_text.split('\n\n')
    for section in sections:
        counter = add_justified_paragraph_with_numbering(section, counter)
        
    #FIGIURES LIST
    elements.append(Paragraph("BRIEF DESCRIPTION OF THE SEVERAL VIEWS OF THE DRAWINGS",heading2_style))
    elements.append(Spacer(1, 12))
    figure_list=Data["list_of_figures"]
    for figure in figure_list:
        counter=add_justified_paragraph_with_numbering(figure,counter)
    
    #Detailed Description   
    elements.append(Paragraph("DETAILED DESCRIPTION",heading2_style))
    elements.append(Spacer(1,12))
     # Method description
    method_desc_text = Data["description"]["method_desc"]["text"]
    method_desc_sections = method_desc_text.split('\n\n')
    for section in method_desc_sections:
        counter =  add_justified_paragraph_with_numbering(section, counter)

    # System description
    system_desc_text_list = Data["description"]["system_desc"]["text_list"]
    for section in system_desc_text_list:
        counter = add_justified_paragraph_with_numbering(section, counter)

    # Invention Description
    invention_desc_text = Data["description"]["invention_desc"]["text"]
    invention_desc_sections = invention_desc_text.split('\n\n')
    for section in invention_desc_sections:
        counter =  add_justified_paragraph_with_numbering(section, counter)
    
    claims_style = ParagraphStyle(
    'Claims',
    parent=style_sheet['BodyText'],
    fontName='Times-Roman',
    fontSize=12,
    leading=20,
    firstLineIndent =40   
)
    # Claims section
    elements.append(PageBreak())
    elements.append(Paragraph("CLAIMS",centered_heading_style))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph("What is claimed is:", ParagraphStyle("font", fontSize=12, fontName="Times-Roman")))
    elements.append(Spacer(1, 12))
    for index, claim in enumerate(Data['claims'], start=1):
        claim_text = f"{index}. {claim['text']}"
    
    # Split the claim text by semicolon
        claim_parts = claim_text.split("\n")
    
    # Add each part of the claim as a new paragraph
        # for part in claim_parts:
        #   elements.append(Paragraph(part, style=claims_style))  # Add ";" back after splitting
        #   elements.append(Spacer(1, 12))
        
        # for i, part in enumerate(claim_parts):
        #     if i < len(claim_parts) - 1:  
        #         elements.append(Paragraph(part.strip() + ";", style=claims_style))
        #     else: 
        #         elements.append(Paragraph(part.strip(), style=claims_style))
        #     elements.append(Spacer(1, 12))
        for part in claim_parts:
            elements.append(Paragraph(part,claims_style))

    # Abstract section
    elements.append(PageBreak())
    elements.append(Paragraph("ABSTRACT", centered_heading_style))
    elements.append(Spacer(1, 12))
    abstract_text = Data["abstract"]["text"]
    abstract_sections = abstract_text.split('\n\n')
    print(abstract_sections)
    for section in abstract_sections:
        counter =  add_justified_paragraph_with_numbering(section,first_Line_Indent=40)
        
    #image section
    elements.append(PageBreak())
    elements.append(Paragraph("FIGURES", heading2_style))
    sorted_claims = sorted(Data["claims"], key=lambda x: x.get("claim_type") == "system")
    for claim in sorted_claims:
        if claim.get("generated_figures_data") and claim["generated_figures_data"].get("latex_details"):
            for latex_detail in claim["generated_figures_data"]["latex_details"]:
                if "images_urls" in latex_detail:
                    for img_url in latex_detail["images_urls"]:
                        img_stream = download_image(img_url)
                        if img_stream:
                            img = Image(img_stream)
                            img_width = img.drawWidth
                            img_height = img.drawHeight 
                            max_width = 5 * inch   
                            max_height = 5 * inch  
                            scale_factor = min(max_width / img_width, max_height / img_height)
                            if scale_factor < 1:
                                img.drawWidth = img_width * scale_factor
                                img.drawHeight = img_height * scale_factor
                            elements.append(img) 
                        else:
                            print(f"Could not download image from {img_url}")

    # i am here building the pdf 
    # pdf.build(elements)
    # Build the PDF with page numbering
    # pdf.build(elements, onFirstPage=add_page_number, onLaterPages=add_page_number)
    pdf.build(elements, onFirstPage=lambda canvas, doc: add_page_number_and_docket(canvas, doc, "Attorney Docket No. "),
              onLaterPages=lambda canvas, doc: add_page_number_and_docket(canvas, doc, "Attorney Docket No. "))
    
    buffer.seek(0)
    return buffer

 

 
