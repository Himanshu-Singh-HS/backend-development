from io import BytesIO
from docx import Document
import docx
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests,json
 
import json

# def load_json(filename):
#     with open(filename, 'r') as file:
#         return json.load(file)

# try:
#     data = load_json("./doc-generator/ss.json")
#     print(data)
# except Exception as e:
#     print(f"An error occurred: {e}")

import json

def load_json(filename):
    try:
        with open(filename, 'r') as file:
            data = file.read()  # Read the file content
            if not data.strip():  # Check if the file is empty or contains only whitespace
                raise ValueError("File is empty or contains only whitespace.")
            return json.loads(data)  # Attempt to parse the JSON data
    except json.JSONDecodeError as e:  # Handle invalid JSON content
        print(f"Invalid JSON format: {e}")
        return None
    except FileNotFoundError as e:  # Handle cases where the file is not found
        print(f"File not found: {e}")
        return None
    except ValueError as e:  # Handle custom empty file error
        print(f"An error occurred: {e}")
        return None
    except Exception as e:  # Catch-all for any other exceptions
        print(f"An unexpected error occurred: {e}")
        return None

# Load the JSON data
data = load_json("./doc-generator/ss.json")
 


class DocGenerator:
    def __init__(self):
        self.doc = Document()

    def add_heading_with_color(self, text, level):
        heading = self.doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        self.doc.add_paragraph()

    def add_paragraph_with_numbering(self, text, counter):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(f"[{counter:04d}] ")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        section_run = paragraph.add_run(text)
        section_run.font.name = 'Times New Roman'
        section_run.font.size = Pt(12)
        
        
            # Add line spacing (e.g., 1.5 lines)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(18)  # 1.5 times the normal spacing (12 * 1.5 = 18)
        paragraph_format.space_after = Pt(12)   #
        
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        return counter + 1  # Increment counter

    def convert_json_to_doc_buffer(self, Data: dict) -> BytesIO:
        """
        Converts a JSON Data to a Word document buffer.

        Args:
            Data (dict): A dictionary containing various sections of the document.
        
        Returns:
            BytesIO: A buffer containing the generated Word document in binary format.
        """
        # Title
        title = Data.get("title", {}).get("text", "Untitled").upper()
        title_heading = self.doc.add_heading(title, level=1)
        title_run = title_heading.runs[0]
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(12)
        # Add line spacing
        title_heading_format = title_heading.paragraph_format
        title_heading_format.line_spacing = Pt(18)  # 1.5 lines
        title_heading_format.space_after = Pt(12)   # 
        
        title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        counter = 1

        # Background attached with technical field
        self.add_heading_with_color('BACKGROUND', level=1)

        # Add technical field text with numbering
        technical_field = Data.get("technical_field", {}).get("text", "")
        if technical_field:
            counter = self.add_paragraph_with_numbering(technical_field, counter)

        # Background text
        background_text = Data.get("background", {}).get("text", "")
        if background_text:
            sections = background_text.split('\n\n')
            for section in sections:
                counter = self.add_paragraph_with_numbering(section, counter)

        # Brief summary text
        self.add_heading_with_color('BRIEF SUMMARY', level=1)
        summary_text = Data.get("summary", {}).get("text", "")
        if summary_text:
            sections = summary_text.split('\n\n')
            for section in sections:
                counter = self.add_paragraph_with_numbering(section, counter)

        # List of figures
        self.add_heading_with_color('BRIEF DESCRIPTION OF THE SEVERAL VIEWS OF THE DRAWINGS', level=1)
        for figure in Data.get("list_of_figures", []):
            counter = self.add_paragraph_with_numbering(figure, counter)

        # Detailed Description
        self.add_heading_with_color('DETAILED DESCRIPTION', level=1)

        # Method description
        method_desc_text = Data.get("description", {}).get("method_desc", {}).get("text", "")
        if method_desc_text:
            method_desc_sections = method_desc_text.split('\n\n')
            for section in method_desc_sections:
                counter = self.add_paragraph_with_numbering(section, counter)

        # System description
        system_desc_text_list = Data.get("description", {}).get("system_desc", {}).get("text_list", [])
        for section in system_desc_text_list:
            cleaned_text = section.replace('\n\n', '')
            counter = self.add_paragraph_with_numbering(cleaned_text, counter)

        # Invention Description
        invention_desc_text = Data.get("description", {}).get("invention_desc", {}).get("text", "")
        if invention_desc_text:
            invention_desc_sections = invention_desc_text.split('\n\n')
            for section in invention_desc_sections:
                counter = self.add_paragraph_with_numbering(section, counter)

        # Claims Description section
        self.doc.add_page_break()
        claims_heading = self.doc.add_heading('CLAIMS', level=1)
        claims_run = claims_heading.runs[0]
        claims_run.font.color.rgb = RGBColor(0, 0, 0)
        claims_run.font.name = 'Times New Roman'
        claims_run.font.size = Pt(12)
        claims_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Claims
        self.doc.add_paragraph("What is claimed is:")
        self.doc.add_paragraph()
        claim_indent = Pt(40)
        for index, claim in enumerate(Data.get('claims', []), start=1):
            claim_parts = claim.get('text', '').split('\n')
            for part in claim_parts:
                claim_paragraph = self.doc.add_paragraph()
                claim_paragraph.paragraph_format.first_line_indent = claim_indent 
                if part == claim_parts[0]:
                    claim_run = claim_paragraph.add_run(f"{index}. ")
                    claim_run.font.name = 'Times New Roman'
                    claim_run.font.size = Pt(12)
                    # Add line spacing
                    
                claim_text_run = claim_paragraph.add_run(part)
                claim_text_run.font.name = 'Times New Roman'
                claim_text_run.font.size = Pt(12)
                # Add line spacing
                title_heading_format =  claim_paragraph.paragraph_format
                title_heading_format.line_spacing = Pt(18)  # 1.5 lines
                title_heading_format.space_after = Pt(12)   # 
            # self.doc.add_paragraph()   

        # Abstract text
        self.doc.add_page_break()
        abstract_heading = self.doc.add_heading('ABSTRACT', level=1)
        self.doc.add_paragraph()
        abstract_run = abstract_heading.runs[0]
        abstract_run.font.color.rgb = RGBColor(0, 0, 0)
        abstract_run.font.name = 'Times New Roman'
        abstract_run.font.size = Pt(12)
        abstract_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        abstract_text = Data.get("abstract", {}).get("text", "")
        abstract_paragraph = self.doc.add_paragraph(abstract_text)

        for run in abstract_paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        abstract_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        abstract_paragraph.paragraph_format.first_line_indent = Pt(50)

        # Figures
        self.doc.add_page_break()
        self.add_heading_with_color('FIGURES', level=1)
        sorted_claims = sorted(Data.get("claims", []), key=lambda x: x.get("claim_type") == "system")

        for claim in sorted_claims:
            generated_figures_data = claim.get("generated_figures_data", {})
            if   generated_figures_data and generated_figures_data.get("latex_details"):
                for latex_detail in generated_figures_data["latex_details"]:
                    if "images_urls" in latex_detail:
                        for img_url in latex_detail["images_urls"]:
                            img_response = requests.get(img_url)
                            if img_response.status_code == 200:
                                image_stream = BytesIO(img_response.content)
                                image_paragraph = self.doc.add_paragraph()
                                run = image_paragraph.add_run()
                                run.add_picture(image_stream, width=Inches(4))
                                image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            else:
                                print(f"Failed to download the image from {img_url}")

        # Add "Attorney Docket No." to the header
        section = self.doc.sections[0]
        header = section.header
        header_paragraph = header.add_paragraph()
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        header_paragraph.add_run("Attorney Docket No.")

        # Function to add page numbering in footer
        def add_page_numbering(doc):
            section = doc.sections[0]
            footer = section.footer
            footer_paragraph = footer.add_paragraph()
            footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = footer_paragraph.add_run()
            run._element.append(
                docx.oxml.parse_xml(
                    r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                )
            )
            run._element.append(
                docx.oxml.parse_xml(
                    r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE  \\* MERGEFORMAT </w:instrText>'
                )
            )
            run._element.append(
                docx.oxml.parse_xml(
                    r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                )
            )

        add_page_numbering(self.doc)

        buffer = BytesIO()
        self.doc.save(buffer)
        self.doc.save("pp.doc")
        print("doc generated sucessfully ")
        buffer.seek(0)
        return buffer

ob=DocGenerator()
ob.convert_json_to_doc_buffer(data)