# from docx import Document
# from docx.shared import RGBColor, Pt, Inches
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from io import BytesIO
# import json

# def json_data(filename):
#     with open(filename, 'r') as file:
#         return json.load(file)

# data = json_data("./EP-doc/ep.json")

# class DocGenerator:
#     def __init__(self):
#         self.doc = Document()
#         self.current_page_number = 1  # Track the current page

#     def add_heading_with_color(self, text, level):
#         heading = self.doc.add_heading(text, level=level)
#         for run in heading.runs:
#             run.font.color.rgb = RGBColor(0, 0, 0)
#             run.font.name = 'Times New Roman'
#             run.font.size = Pt(12)
#         self.doc.add_paragraph()

#     def add_paragraph_with_left_numbering(self, text, number):
#         # Create a table to simulate left-aligned numbering
#         table = self.doc.add_table(rows=1, cols=2)
#         table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#         table.autofit = False
#         table.columns[0].width = Inches(0.3)  # Narrow column for numbering
#         table.columns[1].width = Inches(6.5)  # Wider column for content
        
#         # Add number in the first cell
#         num_cell = table.cell(0, 0)
#         num_paragraph = num_cell.paragraphs[0]
#         num_run = num_paragraph.add_run(f"{number}")
#         num_run.font.size = Pt(10)
#         num_run.font.color.rgb = RGBColor(100, 100, 100)  # Grey for numbering

#         # Add content in the second cell
#         content_cell = table.cell(0, 1)
#         paragraph = content_cell.paragraphs[0]
#         section_run = paragraph.add_run(text)
#         section_run.font.name = 'Times New Roman'
#         section_run.font.size = Pt(12)
#         paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

#     def convert_json_to_doc_buffer(self, Data: dict) -> BytesIO:
#         # Title
#         title = Data.get("title", {}).get("text", "").upper()
#         self.add_heading_with_color(title, level=1)

#         # Background Art Section
#         self.add_heading_with_color('Background Art', level=1)
#         technical_field = Data.get("technical_field", {}).get("text", "")
#         if technical_field:
#             self.add_paragraph_with_left_numbering(technical_field, self.current_page_number)
#             self.current_page_number += 1

#         # Background Content
#         background_text = Data.get("background", {}).get("text", "")
#         if background_text:
#             sections = background_text.split('\n\n')
#             for section in sections:
#                 self.add_paragraph_with_left_numbering(section, self.current_page_number)
#                 self.current_page_number += 1

#         # Add logic for resetting the numbering when a new page starts
#         # This requires splitting content to check each page (usually with manual page breaks)
#         # ...
        
#         # Save document
#         filename = "EPdoc.doc"
#         self.doc.save(filename)
#         print("Document generated successfully ->", filename)

# doc = DocGenerator()
# doc.convert_json_to_doc_buffer(data)


from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json

def json_data(filename):
    with open(filename, 'r') as file:
        return json.load(file)

data = json_data("./EP-doc/ep.json")

class DocGenerator:
    def __init__(self):
        self.doc = Document()

    def add_content_with_margin_numbers(self, text, line_interval=5):
        """Add content with margin numbering at specified intervals."""
        # Table setup for margin numbering
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        table.autofit = False
        table.columns[0].width = Inches(0.3)  # Narrow column for numbers
        table.columns[1].width = Inches(6.5)  # Wider column for content

        # Insert numbers on the left side at the specified intervals
        for i, paragraph_text in enumerate(text.split('\n\n'), start=1):
            if i % line_interval == 1:  # Add number at each interval start
                num_cell = table.cell(0, 0)
                num_paragraph = num_cell.paragraphs[0]
                num_run = num_paragraph.add_run(str(i * line_interval))
                num_run.font.size = Pt(10)
                num_run.font.color.rgb = RGBColor(100, 100, 100)   

            # Add main content in the right cell
            content_cell = table.cell(0, 1)
            content_paragraph = content_cell.add_paragraph(paragraph_text)
            content_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            content_run = content_paragraph.runs[0]
            content_run.font.name = 'Times New Roman'
            content_run.font.size = Pt(12)
            
            # Create a new row for the next paragraph   
            row = table.add_row()

    def convert_json_to_doc_buffer(self, Data: dict):
        # Example usage with title and main content sections
        title = Data.get("title", {}).get("text", "").upper()
        title_heading = self.doc.add_heading(title, level=1)
        title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add main content with margin numbering
     
        
          # Title
        title = Data.get("title", {}).get("text", "").upper()
        title_heading = self.doc.add_heading(title, level=1)
        if title_heading.runs:
            title_run = title_heading.runs[0]
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(12)
        else:
            title_run = title_heading.add_run(title)
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(12)

        title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Background attached with technical field
     

        # Add technical field text with numbering
        technical_field = Data.get("technical_field", {}).get("text", "")
        if technical_field:
            self.add_content_with_margin_numbers(technical_field)

        # Background text
        background_text = Data.get("background", {}).get("text", "")
        if background_text:
            sections = background_text.split('\n\n')
            for section in sections:
                self.add_content_with_margin_numbers(section)

        # Brief summary text
       
        summary_text = Data.get("summary", {}).get("text", "")
        if summary_text:
            sections = summary_text.split('\n\n')
            for section in sections:
                self.add_content_with_margin_numbers(section)

        # List of figures
  
        for figure in Data.get("list_of_figures", []):
            self.add_content_with_margin_numbers(figure)

        

        # Method description
        method_desc_text = Data.get("description", {}).get("method_desc", {}).get("text", "")
        if method_desc_text:
            method_desc_sections = method_desc_text.split('\n\n')
            for section in method_desc_sections:
                self.add_content_with_margin_numbers(section)

        # System description
        system_desc_text_list = Data.get("description", {}).get("system_desc", {}).get("text_list", [])
        if system_desc_text_list:
            for section in system_desc_text_list:
                cleaned_text = section.replace('\n\n', '')
                self.add_content_with_margin_numbers(cleaned_text)
        
        filename = "EPdoc_with_margin_numbers.docx"
        self.doc.save(filename)
        print("Document generated successfully ->", filename)

doc = DocGenerator()
doc.convert_json_to_doc_buffer(data)

