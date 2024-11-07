# from docx import Document

# # Function to check if a line contains non-whitespace text
# def is_non_empty_line(text):
#     # Strip any leading/trailing whitespace and check if there's anything left
#     return bool(text.strip())

# # Path to your Word document
# doc_path = '/Users/patdelanalytics/backend-development/EPdoc.doc'

# # Open and load the Word document
# doc = Document(doc_path)

# # Iterate through each paragraph in the document
# for i, paragraph in enumerate(doc.paragraphs):
#     # Check if the paragraph contains any non-whitespace text
#     if is_non_empty_line(paragraph.text):
#         # Print line number and corresponding text
#         print(f"Line {i+1}: {paragraph.text}")
        
        
from docx import Document
import textwrap

# Function to count words in a string (optional if you want to show word count later)
def is_non_empty_line(text):
    return bool(text.strip())

# Path to your Word document
 
doc_path ='/Users/patdelanalytics/backend-development/EPdoc.doc'

# Open and load the Word document
doc = Document(doc_path)

# Define a maximum number of characters per line (this is a visual approximation)
max_line_length = 90 # Characters per "line" (you can adjust this as per your requirement)

line_number = 1  # Start line numbering

# Iterate through each paragraph in the document
for paragraph in doc.paragraphs:
    # Use textwrap to split the paragraph text into lines based on the max line length
    lines = textwrap.wrap(paragraph.text, width=max_line_length)
    
    # Iterate over each "line" in the wrapped paragraph
    for line in lines:
        if is_non_empty_line(line):  # Only print non-empty lines
            print(f"Line {line_number}: {line}")
            line_number += 1  # Increment line number after printing

