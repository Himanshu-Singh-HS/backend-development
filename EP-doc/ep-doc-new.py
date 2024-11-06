import fitz  # PyMuPDF
from docx import Document

def pdf_to_word_via_text_extraction(pdf_file: str, docx_file: str):
    # Initialize PyMuPDF and a new Word document
    pdf_document = fitz.open(pdf_file)
    word_document = Document()
    
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        text = page.get_text("text")  # Extract plain text from page
        
        # Process and add text to the Word document
        word_document.add_paragraph(text)
        word_document.add_paragraph("\n")  # Add a new line between pages
    
    # Save the Word document
    word_document.save(docx_file)
    pdf_document.close()
    print(f"Extracted text from {pdf_file} to {docx_file}")
    
pdf_to_word_via_text_extraction("/Users/patdelanalytics/backend-development/himanshu.pdf","himanshu.doc")
