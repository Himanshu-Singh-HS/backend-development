
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import json
import docx


def json_data(filename):
    with open (filename,'r') as file:
        return json.load(file)
data=json_data("./EP-doc/ep.json")

class DocGenerator:
    def __init__(self):
        self.doc = Document()
      
        
    def add_heading_with_color(self, text, level):
        heading = self.doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        self.doc.add_paragraph()


    def add_paragraph_with_numbering(self, text):
        paragraph = self.doc.add_paragraph()
        section_run = paragraph.add_run(text)
        section_run.font.name = 'Times New Roman'
        section_run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def convert_json_to_doc_buffer(self, Data: dict) -> BytesIO:

        # Title
        title = Data.get("title", {}).get("text", "").upper()
        title_heading = self.doc.add_heading(title, level=1)
        if title_heading.runs:
            title_run = title_heading.runs[0]
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(12)
        else:
            title_run = title_heading.add_run(title)
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(12)

        title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Background attached with technical field
        self.add_heading_with_color('Background Art', level=1)

        # Add technical field text with numbering
        technical_field = Data.get("technical_field", {}).get("text", "")
        if technical_field:
            self.add_paragraph_with_numbering(technical_field)

        # Background text
        background_text = Data.get("background", {}).get("text", "")
        if background_text:
            sections = background_text.split('\n\n')
            for section in sections:
                self.add_paragraph_with_numbering(section)

        # Brief summary text
        self.add_heading_with_color('Summary of the invention', level=1)
        summary_text = Data.get("summary", {}).get("text", "")
        if summary_text:
            sections = summary_text.split('\n\n')
            for section in sections:
                self.add_paragraph_with_numbering(section)

        # List of figures
        self.add_heading_with_color('Brief description of the drawings', level=1)
        for figure in Data.get("list_of_figures", []):
            self.add_paragraph_with_numbering(figure)

        # Detailed Description
        self.add_heading_with_color('Detailed description', level=1)

        # Method description
        method_desc_text = Data.get("description", {}).get("method_desc", {}).get("text", "")
        if method_desc_text:
            method_desc_sections = method_desc_text.split('\n\n')
            for section in method_desc_sections:
                self.add_paragraph_with_numbering(section)

        # System description
        system_desc_text_list = Data.get("description", {}).get("system_desc", {}).get("text_list", [])
        if system_desc_text_list:
            for section in system_desc_text_list:
                cleaned_text = section.replace('\n\n', '')
                self.add_paragraph_with_numbering(cleaned_text)

        # Invention Description
        invention_desc_text = Data.get("description", {}).get("invention_desc", {}).get("text", "")
        if invention_desc_text:
            invention_desc_sections = invention_desc_text.split('\n\n')
            for section in invention_desc_sections:
                self.add_paragraph_with_numbering(section)

        # Claims Description section
        self.doc.add_page_break()
        claims_heading = self.doc.add_heading('Claims', level=1)
        self.doc.add_paragraph()
        claims_run = claims_heading.runs[0]
        claims_run.font.color.rgb = RGBColor(0, 0, 0)
        claims_run.font.name = 'Times New Roman'
        claims_run.font.size = Pt(12)
        
            
        for index, claim in enumerate(Data.get("claims", []), start=1):
            claim_parts = claim.get('text', '').split('\n')
            for part in claim_parts:
                claim_paragraph = self.doc.add_paragraph()
                if part == claim_parts[0]:
                    claim_run = claim_paragraph.add_run(f"{index}. ")
                    claim_run.font.name = 'Times New Roman'
                    claim_run.font.size = Pt(12)
                    claim_paragraph.add_run(' ' * 6)
                claim_text_run = claim_paragraph.add_run(part)
                claim_text_run.font.name = 'Times New Roman'
                claim_text_run.font.size = Pt(12)
            
            
        # Abstract text
        self.doc.add_page_break()
        abstract_heading = self.doc.add_heading('Abstract', level=1)
        self.doc.add_paragraph()
        abstract_run = abstract_heading.runs[0]
        abstract_run.font.color.rgb = RGBColor(0, 0, 0)
        abstract_run.font.name = 'Times New Roman'
        abstract_run.font.size = Pt(12)

        abstract_text = Data.get("abstract", {}).get("text", "")
        abstract_paragraph = self.doc.add_paragraph(abstract_text)
        for run in abstract_paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        abstract_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Function to add page numbering in footer
        def add_page_numbering(doc):
            section = doc.sections[0]
            header = section.header
            header.margin_top = Pt(20)
            header_paragraph = header.add_paragraph()
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = header_paragraph.add_run()
            run._element.append(
                docx.oxml.parse_xml(
                    r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                )
            )
            run._element.append(
                docx.oxml.parse_xml(
                    r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE  \\* MERGEFORMAT </w:instrText>'
                )
            )
            run._element.append(
                docx.oxml.parse_xml(
                    r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                )
            )

        add_page_numbering(self.doc)
        filename="EPdoc.doc"
        self.doc.save(filename)
        print("Document generated successfully->",filename)

doc = DocGenerator()
doc.convert_json_to_doc_buffer(data)
