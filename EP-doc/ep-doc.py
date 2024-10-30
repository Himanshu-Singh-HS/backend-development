import json
from io import BytesIO
import logging
from docx import Document
import docx
from docx.shared import Pt, Inches,Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.shared import RGBColor
import requests,time
from docx.enum.text import WD_ALIGN_PARAGRAPH

def json_data(filename):
    with open (filename,'r') as file:
        return json.load(file)
data=json_data("./EP-doc/ep.json")
 


class DocGenerator:
    
    def __init__(self):
        self.doc = Document()
    print("*********************generating doc----------------------")
    
    
    #code added for EP
    
    def add_heading_with_color(self,text, level):
            heading = self.doc.add_paragraph(text)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
            heading.paragraph_format.space_after = Pt(0)
            
    def add_paragraph_with_numbering(self,text, counter):
            paragraph = self.doc.add_paragraph()
            paragraph.paragraph_format.line_spacing = Pt(18)  
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(f"[{counter:04d}] ")
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            paragraph.add_run(' ' * 4)
            section_run = paragraph.add_run(text)
            section_run.font.name = 'Times New Roman'
            section_run.font.size = Pt(10)
            return counter + 1   
        
    def convert_json_to_doc_buffer(self,Data: dict):
        # Title
        title = Data.get("title", {}).get("text", "").upper()
        title_heading = self.doc.add_heading(title, level=1)
        if title_heading.runs:
            title_run = title_heading.runs[0]
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(10)
        else:
            title_run = title_heading.add_run(title)
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(10)
        title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        
        counter = 1
        
        # Background attached with technical field
        self.add_heading_with_color('Background Art', level=2)

        # Add technical field text with numbering
        technical_field = Data.get("technical_field", {}).get("text", "")
        if technical_field:
            counter = self.add_paragraph_with_numbering(technical_field, counter)

        # Background text
        background_text = Data.get("background", {}).get("text", "")
        if background_text:
            sections = background_text.split('\n\n')
            for section in sections:
                counter = self.add_paragraph_with_numbering(section, counter)

        # Brief summary text
        self.add_heading_with_color('Summary of the invention', level=1)
        summary_text=  Data.get("summary", {}).get("text", "")
        if summary_text:
            sections=summary_text.split('\n\n')
            for section in sections:
                counter = self.add_paragraph_with_numbering(section, counter)

        # List of figures
        self.doc.add_page_break()
        self.add_heading_with_color(
            'Brief description of the drawings', level=1)
        for figure in  Data.get("list_of_figures", []):
            counter = self.add_paragraph_with_numbering(figure, counter)

        # Detailed Description
        self.doc.add_page_break()
        self.add_heading_with_color('Detailed description', level=1)

        # Method description
        method_desc_text =  Data.get("description", {}).get("method_desc", {}).get("text", "")
        if  method_desc_text:
            method_desc_sections = method_desc_text.split('\n\n')
            for section in method_desc_sections:
                counter = self.add_paragraph_with_numbering(section, counter)

        # System description
        system_desc_text_list = Data.get("description", {}).get("system_desc", {}).get("text_list", [])
        if system_desc_text_list:
            for section in system_desc_text_list:
                cleaned_text = section.replace('\n\n', '')
                counter = self.add_paragraph_with_numbering(cleaned_text, counter)

        # Invention Description
        invention_desc_text =  Data.get("description", {}).get("invention_desc", {}).get("text", "")
        if  invention_desc_text:
            invention_desc_sections = invention_desc_text.split('\n\n')
            for section in invention_desc_sections:
                counter = self.add_paragraph_with_numbering(section, counter)

        # Claims section
        self.doc.add_page_break()
        claims_heading = self.doc.add_paragraph("Claims")
        claims_run = claims_heading.runs[0]
        claims_run.font.color.rgb = RGBColor(0, 0, 0)
        claims_run.font.name = 'Times New Roman'
        claims_run.font.size = Pt(10)
        
        # # Claims
        # for index, claim in enumerate(Data.get("claims", []), start=1):
        #     claim_parts = claim.get('text', '').split('\n')
        #     for part in claim_parts:
        #         claim_paragraph = self.doc.add_paragraph()
        #         if part == claim_parts[0]:
        #             claim_run = claim_paragraph.add_run(f"{index}. ")
        #             claim_run.font.name = 'Times New Roman'
        #             claim_run.font.size = Pt(10)
        #             claim_paragraph.add_run(' ' * 6)
        #         claim_text_run = claim_paragraph.add_run(part)
        #         claim_text_run.font.name = 'Times New Roman'
        #         claim_text_run.font.size = Pt(10)
        
        # # Claims
        # line_count = 0  # Initialize line count
        # for index, claim in enumerate(Data.get("claims", []), start=1):
        #     claim_parts = claim.get('text', '').split('\n')
        #     for part in claim_parts:
        #         claim_paragraph = self.doc.add_paragraph()
                
        #         # Increment line count for every part
        #         line_count += 1
                
        #         # Add left-side count for every line
        #         if line_count % 5 == 0:
        #             line_number = str(line_count)
        #             line_number_run = claim_paragraph.add_run(f"{line_number} ")
        #             line_number_run.font.name = 'Times New Roman'
        #             line_number_run.font.size = Pt(10)

        #         # Add the claim index number for the first part
        #         if part == claim_parts[0]:
        #             claim_run = claim_paragraph.add_run(f"{index}. ")
        #             claim_run.font.name = 'Times New Roman'
        #             claim_run.font.size = Pt(10)

        #         # Add some space after the index
        #         claim_paragraph.add_run(' ' * 4)

        #         # Add the claim text
        #         claim_text_run = claim_paragraph.add_run(part)
        #         claim_text_run.font.name = 'Times New Roman'
        #         claim_text_run.font.size = Pt(10)
        
        # Claims
        # line_count = 0   
        # for index, claim in enumerate(Data.get("claims", []), start=1):
        #     claim_parts = claim.get('text', '').split('\n')
        #     for part in claim_parts:
        #         line_count += 1
        #         claim_paragraph = self.doc.add_paragraph()

        #         # Handle numbering for every 5th line (left-side line numbers)
        #         if line_count % 5 == 0:
        #             line_number_paragraph = self.doc.add_paragraph()
        #             line_number = str(line_count)
        #             line_number_run = line_number_paragraph.add_run(f"{line_number} ")
        #             line_number_run.font.name = 'Times New Roman'
        #             line_number_run.font.size = Pt(10)
        #             # Align the line numbers to the left outside
        #             line_number_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        #         if part == claim_parts[0]:
        #             claim_run = claim_paragraph.add_run(f"{index}. ")
        #             claim_run.font.name = 'Times New Roman'
        #             claim_run.font.size = Pt(10)
        #         claim_paragraph.add_run(' ' * 4)
                
        #         # Add the claim text
        #         claim_text_run = claim_paragraph.add_run(part)
        #         claim_text_run.font.name = 'Times New Roman'
        #         claim_text_run.font.size = Pt(10)
        line_count = 0   
        for index, claim in enumerate(Data.get("claims", []), start=1):
            claim_parts = claim.get('text', '').split('\n')
            for part in claim_parts:
                line_count += 1

                # Create a new paragraph for each claim part
                claim_paragraph = self.doc.add_paragraph()

                # Add line number every 5th line outside of the content
                if line_count % 5 == 1:  # Numbering starts at the beginning of every fifth set
                    line_number = str(line_count)
                    claim_paragraph.add_run(f"{line_number} ").font.size = Pt(10)

                # Add the claim number for the first line of the claim part
                if part == claim_parts[0]:
                    claim_run = claim_paragraph.add_run(f"{index}. ")
                    claim_run.font.name = 'Times New Roman'
                    claim_run.font.size = Pt(10)

                # Add the claim text without additional spacing
                claim_text_run = claim_paragraph.add_run(part)
                claim_text_run.font.name = 'Times New Roman'
                claim_text_run.font.size = Pt(10)


 





        
        # =============================================================================
        # Abstract text
        self.doc.add_page_break()
        abstract_heading = self.doc.add_paragraph("ABSTRACT")
        abstract_run = abstract_heading.runs[0]
        abstract_run.font.color.rgb = RGBColor(0, 0, 0)
        abstract_run.font.name = 'Times New Roman'
        abstract_run.font.size = Pt(10)
        
        abstract_text = Data.get("abstract", {}).get("text", "")
        abstract_paragraph = self.doc.add_paragraph(abstract_text)

        for run in abstract_paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

       
        abstract_paragraph.paragraph_format.first_line_indent = Pt(50)

        # Figures
        self.doc.add_page_break()
        self.add_heading_with_color('FIGURES', level=1)
        sorted_claims = sorted(Data.get("claims", []), key=lambda x: x.get("claim_type") == "system")

        for claim in sorted_claims:
            if claim.get("generated_figures_data",{}) and claim["generated_figures_data"].get("latex_details"):
                for latex_detail in claim["generated_figures_data"]["latex_details"]:
                    if "images_urls" in latex_detail:
                        for img_url in latex_detail["images_urls"]:
                            try:
                                img_response = requests.get(img_url)
                                if img_response.status_code == 200:
                                    image_stream = BytesIO(img_response.content)
                                    image_paragraph = self.doc.add_paragraph()
                                    run = image_paragraph.add_run()
                                    run.add_picture(image_stream, width=Inches(4))
                                    image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            except requests.exceptions.RequestException as e:
                                  logging.error(f"Failed to download the image from {img_url}: {e}")
                            except Exception as e:
                                   logging.error(f"Unexpected error while processing image from {img_url}: {e}")

        # # Add "Attorney Docket No." to the header
        # section = self.doc.sections[0]
        # header = section.header
        # header_paragraph = header.add_paragraph()
        # header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        # header_paragraph.add_run("Attorney Docket No.")

        # Function to add page numbering in header
        def add_page_numbering(doc):
            section = doc.sections[0]
            header = section.header
            header.margin_top = Pt(20) #added by EP
            header_paragraph = header.add_paragraph()
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # header_paragraph.paragraph_format.space_after = Pt(10) #added by EP
            run = header_paragraph.add_run()
            run._element.append(
                docx.oxml.parse_xml(
                    r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                )
            )
            run._element.append(
                docx.oxml.parse_xml(
                    r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE  \\* MERGEFORMAT </w:instrText>'
                )
            )
            run._element.append(
                docx.oxml.parse_xml(
                    r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                )
            )

        add_page_numbering(self.doc)
        self.doc.save("EP.doc")
        print("doc generated sucessfully ")
        


doc=DocGenerator()
doc.convert_json_to_doc_buffer(data)