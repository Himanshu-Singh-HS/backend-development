from docx import Document

# Function to count words in a string
def count_words_in_text(text):
    # Split the text into words using whitespace
    words = text.split()
    return len(words)

doc_path = '/Users/patdelanalytics/backend-development/EPdoc.doc'
doc = Document(doc_path)

# # Initialize a dictionary to store word counts per line
# line_word_counts = {}

# # Iterate through each paragraph in the document (paragrahps are generally line-like in Word documents)
# for i, paragraph in enumerate(doc.paragraphs):
#     # Count the words in the current paragraph
#     word_count = count_words_in_text(paragraph.text)
    
#     # Store the word count for this paragraph (or line-like structure)
#     line_word_counts[f'Line {i+1}'] = paragraph.text

# for line, word_count in line_word_counts.items():  
#     print(f"{line}: {word_count}")


# Iterate through each paragraph in the document
for i, paragraph in enumerate(doc.paragraphs):
    # Print line number and corresponding text (paragraph content)
    print(f"Line {i+1}: {paragraph.text}")
