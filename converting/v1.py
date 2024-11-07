import PyPDF2
from docx import Document

# Extract text from PDF
pdf_file = '/Users/patdelanalytics/backend-development/converting/xyz.pdf'
docx_file = 'output92.doc'

# with open(pdf_file, 'rb') as file:
#     reader = PyPDF2.PdfReader(file)
#     document = Document()

#     for page_num in range(len(reader.pages)):
#         page = reader.pages[page_num]
#         text = page.extract_text()
#         document.add_paragraph(text)

#     # Save as Word Document
#     document.save(docx_file)

# print("Conversion Complete!")


#  80percent working 
from pdf2docx import Converter

# Convert the PDF to DOCX
docx_file = 'output01.docx'

# Create a converter object and convert
cv = Converter(pdf_file)
cv.convert(docx_file, start=0, end=None,layout=True, image=True)  # start and end are page numbers
cv.close()

print("Conversion Complete!")






# not working correctly 
# import fitz  # PyMuPDF
# from docx import Document

# docx_file = 'output9.docx'

# # Open the PDF file
# doc = fitz.open(pdf_file)
# document = Document()

# # Loop through each page
# for page_num in range(len(doc)):
#     page = doc.load_page(page_num)
#     text = page.get_text("text")  # Extract text

#     # Add extracted text as a new paragraph
#     document.add_paragraph(text)

# # Save to DOCX
# document.save(docx_file)

# print("Conversion Complete!")

