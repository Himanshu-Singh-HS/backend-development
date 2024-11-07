from docx import Document
from docx.shared import Cm

def add_line_numbers_with_page_reset(doc_path, output_path, max_lines_per_page=40):
    """
    Adds line numbers in the left margin area for each line and resets numbering on each page.
    """
    # Load the original document
    doc = Document(doc_path)
    new_doc = Document()  # Create a new document for the formatted output

    line_number = 1  # Start line numbering
    lines_on_current_page = 0  # Track lines per page

    # Iterate over each paragraph in the document
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Only process non-empty paragraphs
            original_text = paragraph.text
            line_texts = [original_text[i:i + 80] for i in range(0, len(original_text), 80)]  # Split into lines (approx.)

            # Create a two-column table for each line in the paragraph
            for line_text in line_texts:
                # Create a table with two columns: line numbers and content
                table = new_doc.add_table(rows=1, cols=2)
                table.autofit = False
                table.columns[0].width = Cm(1.5)  # Width for line numbers column
                table.columns[1].width = Cm(14)   # Width for content column

                # Add line number in the left cell
                line_number_cell = table.cell(0, 0)
                line_number_cell.text = f"{line_number}."

                # Add line text in the right cell, preserving original paragraph style
                content_cell = table.cell(0, 1)
                content_paragraph = content_cell.paragraphs[0]
                content_paragraph.style = paragraph.style  # Preserve original formatting
                content_paragraph.text = line_text

                # Preserve the original alignment (justification)
                content_paragraph.alignment = paragraph.alignment

                # Ensure content matches original text format (font, size, color, etc.)
                for run in content_paragraph.runs:
                    # Ensure each run preserves the formatting from the original paragraph
                    for original_run in paragraph.runs:
                        if run.text == original_run.text:
                            run.font.name = original_run.font.name
                            run.font.size = original_run.font.size
                            run.font.bold = original_run.font.bold
                            run.font.italic = original_run.font.italic
                            run.font.underline = original_run.font.underline
                            run.font.color.rgb = original_run.font.color.rgb

                # Increment line number and count lines on the current page
                line_number += 1
                lines_on_current_page += 1

                # Reset line numbers if we reach the max lines per page
                if lines_on_current_page >= max_lines_per_page:
                    line_number = 1  # Reset line number for new page
                    lines_on_current_page = 0  # Reset page line count

    # Save the new document with line numbers on the left side
    new_doc.save(output_path)
    print(f"Document saved with line numbers that reset on each page at {output_path}")

 

# Example usage
add_line_numbers_with_page_reset("EPdoc.doc", "qqqq.doc")
