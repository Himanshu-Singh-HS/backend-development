# from docx import Document
# from docx.shared import Pt, Cm

# def add_line_numbers_by_line(doc_path, output_path, max_chars_per_line=90, max_lines_per_page=35):
#     """
#     Add line numbers line by line, restarting on each page.
#     """
#     # Load the document
#     doc = Document(doc_path)
    
#     line_number = 1
#     lines_on_current_page = 0

#     # Iterate over each paragraph
#     for paragraph in doc.paragraphs:
#         # Only proceed if the paragraph has text
#         if paragraph.text.strip():
#             # Clear the original text to insert line-by-line text
#             original_text = paragraph.text
#             paragraph.clear()

#             # Split paragraph text into lines based on character count per line
#             for i in range(0, len(original_text), max_chars_per_line):
#                 line_text = original_text[i:i+max_chars_per_line]
                
#                 # Add line number with the line text
#                 line = f"{line_number}. {line_text}"
#                 run = paragraph.add_run(line)
#                 run.font.size = Pt(11)  # Set font size as needed

#                 # Increment line number and page line counter
#                 line_number += 1
#                 lines_on_current_page += 1

#                 # Check if we need to reset line number for a new "page"
#                 if lines_on_current_page >= max_lines_per_page:
#                     line_number = 1
#                     lines_on_current_page = 0

#             # Set paragraph formatting
#             paragraph.paragraph_format.left_indent = Cm(1.5)  # Adjust for left margin space
#             paragraph.paragraph_format.line_spacing = Pt(12)  # Set line spacing

#     # Save the modified document
#     doc.save(output_path)
#     print(f"Document saved with line-by-line numbering at {output_path}")

# Example usage


# from docx import Document
# from docx.shared import Pt, Cm

# def add_line_numbers_left_side(doc_path, output_path, max_chars_per_line=80, max_lines_per_page=40):
#     """
#     Adds line numbers in a left margin-like column for each line of content in a Word document.
#     """
#     # Load the document
#     doc = Document(doc_path)
#     new_doc = Document()  # Create a new document to save formatted content

#     line_number = 1
#     lines_on_current_page = 0

#     # Iterate over each paragraph
#     for paragraph in doc.paragraphs:
#         # Only proceed if the paragraph has text
#         if paragraph.text.strip():
#             # Split paragraph text into lines based on character count per line
#             original_text = paragraph.text
#             for i in range(0, len(original_text), max_chars_per_line):
#                 line_text = original_text[i:i+max_chars_per_line]

#                 # Create a table with two columns: one for line numbers, one for content
#                 table = new_doc.add_table(rows=1, cols=2)
#                 table.autofit = False
#                 table.columns[0].width = Cm(1.5)  # Adjust the width of the line number column
#                 table.columns[1].width = Cm(14)  # Adjust the width of the content column

#                 # Add line number in the left cell
#                 line_number_cell = table.cell(0, 0)
#                 line_number_cell.text = f"{line_number}."
#                 line_number_cell.paragraphs[0].runs[0].font.size = Pt(11)  # Set font size as needed

#                 # Add line text in the right cell
#                 content_cell = table.cell(0, 1)
#                 content_cell.text = line_text
#                 content_cell.paragraphs[0].runs[0].font.size = Pt(11)  # Set font size as needed

#                 # Increment line number and page line counter
#                 line_number += 1
#                 lines_on_current_page += 1

#                 # Check if we need to reset line number for a new "page"
#                 if lines_on_current_page >= max_lines_per_page:
#                     line_number = 1
#                     lines_on_current_page = 0

#     # Save the new document with line numbers on the left side
#     new_doc.save(output_path)
#     print(f"Document saved with line numbers in the left margin at {output_path}")

# # Example usage
# add_line_numbers_left_side("EPdoc.doc", "lllll.docx")


from docx import Document
from docx.shared import Cm

def add_line_numbers_left_side_preserve_formatting(doc_path, output_path, max_chars_per_line=80, max_lines_per_page=40):
    """
    Add line numbers in the left margin area of each line without altering original text formatting.
    """
    # Load the original document
    doc = Document(doc_path)
    new_doc = Document()  # Create a new document to add formatted content

    line_number = 1
    lines_on_current_page = 0

    # Iterate over each paragraph
    for paragraph in doc.paragraphs:
        # Only proceed if the paragraph has text
        if paragraph.text.strip():
            original_text = paragraph.text
            # Split the paragraph into lines based on character count
            for i in range(0, len(original_text), max_chars_per_line):
                line_text = original_text[i:i+max_chars_per_line]

                # Create a two-column table for each line
                table = new_doc.add_table(rows=1, cols=2)
                table.autofit = False
                table.columns[0].width = Cm(1.5)  # Width for line numbers column
                table.columns[1].width = Cm(14)  # Width for content column

                # Add line number in the left cell
                line_number_cell = table.cell(0, 0)
                line_number_cell.text = f"{line_number}."

                # Add line text in the right cell and apply original formatting
                content_cell = table.cell(0, 1)
                content_cell.text = line_text

                # Copy paragraph formatting to preserve style
                new_paragraph = content_cell.paragraphs[0]
                new_paragraph.style = paragraph.style

                # Increment line number and check for page reset
                line_number += 1
                lines_on_current_page += 1

                if lines_on_current_page >= max_lines_per_page:
                    line_number = 1
                    lines_on_current_page = 0

    # Save the modified document
    new_doc.save(output_path)
    print(f"Document saved with line numbers on the left side at {output_path}")

# Example usage
add_line_numbers_left_side_preserve_formatting("EPdoc.doc", "output.docx")
